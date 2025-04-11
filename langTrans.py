from tkinter import *
from tkinter import filedialog, messagebox, scrolledtext
from PIL import Image
import pytesseract
import pdfplumber
import speech_recognition as sr
from googletrans import Translator
from docx import Document
from gtts import gTTS

# Initialize translator
translator = Translator()

# Defining the process functions
def upload_file():
    file_type = file_type_var.get()
    if file_type == "Image":
        process_image()
    elif file_type == "Audio":
        process_audio()
    elif file_type == "PDF":
        process_pdf()
    elif file_type == "DOCX":
        process_docx()

def process_image():
    file_path = filedialog.askopenfilename(filetypes=[("Image files", "*.png;*.jpg;*.jpeg")])
    if not file_path:
        return
    img = Image.open(file_path)
    text = pytesseract.image_to_string(img)
    input_textbox.delete('1.0', END)
    input_textbox.insert(END, text)

def process_audio():
    file_path = filedialog.askopenfilename(filetypes=[("Audio files", "*.wav;*.aiff;*.flac")])
    if not file_path:
        return
    recognizer = sr.Recognizer()
    with sr.AudioFile(file_path) as source:
        audio = recognizer.record(source)
    try:
        text = recognizer.recognize_google(audio)
    except sr.UnknownValueError:
        text = "Could not understand audio"
    except sr.RequestError:
        text = "Could not request results; check your network connection"
    input_textbox.delete('1.0', END)
    input_textbox.insert(END, text)

def process_pdf():
    file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
    if not file_path:
        return
    text = ''
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text()
    input_textbox.delete('1.0', END)
    input_textbox.insert(END, text)

def process_docx():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if not file_path:
        return
    doc = Document(file_path)
    text = ''
    for para in doc.paragraphs:
        text += para.text
    input_textbox.delete('1.0', END)
    input_textbox.insert(END, text)

def translate_text():
    src_text = input_textbox.get('1.0', END).strip()
    dest_lang = languages[language_var.get()]
    if not src_text:
        messagebox.showerror("Input Error", "Please provide some text to translate.")
        return
    translated = translator.translate(src_text, dest=dest_lang)
    output_textbox.delete('1.0', END)
    output_textbox.insert(END, translated.text)

def save_as_docx(text, file_path):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(file_path)

def save_as_audio(text, file_path):
    language_code = languages[language_var.get()]  # Get the selected language code
    tts = gTTS(text=text, lang=language_code)
    tts.save(file_path)

def save_output():
    text = output_textbox.get('1.0', END).strip()
    if not text:
        messagebox.showerror("Input Error", "There's no text to save.")
        return

    file_types = [("DOCX file", "*.docx"), ("Audio file", "*.mp3")]
    file_path = filedialog.asksaveasfilename(filetypes=file_types, defaultextension=file_types)

    if file_path.endswith('.docx'):
        save_as_docx(text, file_path)
    elif file_path.endswith('.mp3'):
        save_as_audio(text, file_path)
    else:
        messagebox.showerror("File Type Error", "Unsupported file type selected.")

# Create the main window
window = Tk()
window.title("TranslateX")
window.geometry("1200x700")
window.minsize(1000, 600)
window.config(bg="#E6F2FA")  # Light blue background

# Styling
colors = {
    "light_blue": "#A7C7E7",
    "dark_blue": "#4A90E2",
    "neutral": "#F5F5F5",
    "text_bg": "#FFFFFF",
    "text_fg": "#000000",
    "frame_bg": "#FFFFFF",
    "light_green": "#A8D5BA",
    "cool_green": "#88C0A8"
}

# Hover effect for buttons
def on_enter(e, color):
    e.widget['background'] = color

def on_leave(e, color):
    e.widget['background'] = color

# Create frames for input and output sections
input_frame = Frame(window, bg=colors["frame_bg"], padx=10, pady=10, relief=RIDGE, bd=2)
input_frame.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")

output_frame = Frame(window, bg=colors["frame_bg"], padx=10, pady=10, relief=RIDGE, bd=2)
output_frame.grid(row=0, column=1, padx=10, pady=10, sticky="nsew")

# Create a sub-frame for the upload section
upload_frame = Frame(input_frame, bg=colors["frame_bg"])
upload_frame.pack(pady=5)

# Create dropdown menu for file type selection in upload frame
file_type_var = StringVar(value='Image')
file_types = ['Image', 'Audio', 'PDF', 'DOCX']
file_type_menu = OptionMenu(upload_frame, file_type_var, *file_types)
file_type_menu.grid(row=0, column=0, padx=5)

# Create the upload button next to the dropdown menu
upload_button = Button(upload_frame, text="Upload File", command=upload_file, bg=colors["dark_blue"], fg=colors["text_fg"], relief=RAISED)
upload_button.grid(row=0, column=1, padx=5)
upload_button.bind("<Enter>", lambda e: on_enter(e, colors["light_blue"]))
upload_button.bind("<Leave>", lambda e: on_leave(e, colors["dark_blue"]))

# Create the input textbox
input_textbox = scrolledtext.ScrolledText(input_frame, wrap=WORD, height=20, width=60, bg=colors["text_bg"], fg=colors["text_fg"], relief=SUNKEN, bd=2)
input_textbox.pack(padx=10, pady=10, fill=BOTH, expand=True)

# Create a centered frame for language selection
language_frame = Frame(output_frame, bg=colors["frame_bg"])
language_frame.pack(pady=10)

# Create dropdown menu for language selection on output frame with reduced width
language_var = StringVar(value='Hindi')
languages = {'Hindi': 'hi', 'Marathi': 'mr', 'Bengali': 'bn', 'Gujarati': 'gu', 'Tamil': 'ta', 'Telugu': 'te'}
language_menu = OptionMenu(language_frame, language_var, *languages.keys())
language_menu.config(width=10, bg=colors["neutral"])
language_menu.pack(anchor=CENTER)

# Create the output textbox
output_textbox = scrolledtext.ScrolledText(output_frame, wrap=WORD, height=20, width=60, bg=colors["text_bg"], fg=colors["text_fg"], relief=SUNKEN, bd=2)
output_textbox.pack(padx=10, pady=10, fill=BOTH, expand=True)

# Create the process button with light green color
process_button = Button(window, text="Translate Text", command=translate_text, bg=colors["light_green"], fg=colors["text_fg"], relief=RAISED)
process_button.grid(row=1, column=0, pady=10, padx=10, sticky="ew")
process_button.bind("<Enter>", lambda e: on_enter(e, "#B7E4CB"))
process_button.bind("<Leave>", lambda e: on_leave(e, colors["light_green"]))

# Create the save button with a matching cool green color
save_button = Button(window, text="Save Output", command=save_output, bg=colors["cool_green"], fg=colors["text_fg"], relief=RAISED)
save_button.grid(row=1, column=1, pady=10, padx=10, sticky="ew")
save_button.bind("<Enter>", lambda e: on_enter(e, "#9FDCC1"))
save_button.bind("<Leave>", lambda e: on_leave(e, colors["cool_green"]))

# Configure window grid layout
window.grid_rowconfigure(0, weight=1)
window.grid_columnconfigure(0, weight=1)
window.grid_columnconfigure(1, weight=1)

window.mainloop()
